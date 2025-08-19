#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PROCESSEUR DE DOCUMENTS - RAG DASHBOARD MALAYSIA
===============================================

Système pour traiter et indexer des documents (PDF, Word, TXT, etc.)
dans la base de connaissances RAG pour enrichir les réponses du LLM

Version: 1.0.0
"""

import os
import logging
import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import sqlite3
import json

import pandas as pd
import numpy as np

# Processors par type de fichier
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processeur de documents pour l'alimentation RAG"""
    
    def __init__(self, rag_service, documents_dir: str = "documents"):
        """
        Initialise le processeur de documents
        
        Args:
            rag_service: Instance du service RAG
            documents_dir: Dossier contenant les documents à traiter
        """
        self.rag_service = rag_service
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(exist_ok=True)
        
        # Dossiers par type
        self.processed_dir = self.documents_dir / "processed"
        self.failed_dir = self.documents_dir / "failed"
        self.processed_dir.mkdir(exist_ok=True)
        self.failed_dir.mkdir(exist_ok=True)
        
        # Types de fichiers supportés
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.csv': self._process_csv,
            '.json': self._process_json,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel
        }
        
        # Configuration du chunking (découpage en morceaux)
        self.chunk_config = {
            'max_chunk_size': 1000,  # Caractères par chunk
            'overlap_size': 100,     # Chevauchement entre chunks
            'min_chunk_size': 50     # Taille minimale d'un chunk
        }
        
        # Base de données des documents traités
        self._init_documents_db()
        
        logger.info("✅ DocumentProcessor initialisé")
    
    def _init_documents_db(self):
        """Initialise la base de données des documents"""
        db_path = self.documents_dir / "documents_index.db"
        
        with sqlite3.connect(db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS processed_documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    filename TEXT UNIQUE,
                    file_path TEXT,
                    file_hash TEXT,
                    file_size INTEGER,
                    file_type TEXT,
                    processing_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    chunks_created INTEGER,
                    processing_status TEXT,
                    error_message TEXT,
                    metadata TEXT
                )
            """)
            
            conn.execute("""
                CREATE TABLE IF NOT EXISTS document_chunks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    document_id INTEGER,
                    chunk_index INTEGER,
                    content TEXT,
                    content_hash TEXT UNIQUE,
                    metadata TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (document_id) REFERENCES processed_documents (id)
                )
            """)
            
            conn.commit()
        
        self.db_path = db_path
    
    def process_documents_directory(self) -> Dict:
        """
        Traite tous les documents dans le dossier documents/
        
        Returns:
            Dict: Rapport de traitement
        """
        try:
            logger.info(f"🔍 Scan du dossier: {self.documents_dir}")
            
            # Recherche de tous les fichiers
            all_files = []
            for file_path in self.documents_dir.rglob('*'):
                if file_path.is_file() and not file_path.name.startswith('.'):
                    # Éviter les dossiers de traitement
                    if 'processed' not in str(file_path) and 'failed' not in str(file_path):
                        all_files.append(file_path)
            
            logger.info(f"📁 {len(all_files)} fichiers trouvés")
            
            # Traitement des fichiers
            results = {
                'total_files': len(all_files),
                'processed': 0,
                'skipped': 0,
                'failed': 0,
                'new_chunks': 0,
                'processing_details': []
            }
            
            for file_path in all_files:
                try:
                    result = self.process_single_document(file_path)
                    
                    if result['success']:
                        results['processed'] += 1
                        results['new_chunks'] += result.get('chunks_created', 0)
                    elif result.get('skipped'):
                        results['skipped'] += 1
                    else:
                        results['failed'] += 1
                    
                    results['processing_details'].append({
                        'file': str(file_path.name),
                        'status': 'success' if result['success'] else 'failed',
                        'chunks': result.get('chunks_created', 0),
                        'message': result.get('message', '')
                    })
                    
                except Exception as e:
                    logger.error(f"❌ Erreur traitement {file_path.name}: {e}")
                    results['failed'] += 1
                    results['processing_details'].append({
                        'file': str(file_path.name),
                        'status': 'error',
                        'chunks': 0,
                        'message': str(e)
                    })
            
            logger.info(f"✅ Traitement terminé: {results['processed']} réussis, {results['failed']} échecs")
            return results
            
        except Exception as e:
            logger.error(f"❌ Erreur traitement dossier: {e}")
            return {
                'total_files': 0,
                'processed': 0,
                'failed': 0,
                'error': str(e)
            }
    
    def process_single_document(self, file_path: Path) -> Dict:
        """
        Traite un document unique
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Dict: Résultat du traitement
        """
        try:
            # Vérification de l'existence
            if not file_path.exists():
                return {'success': False, 'message': 'Fichier introuvable'}
            
            # Calcul du hash pour détecter les doublons
            file_hash = self._calculate_file_hash(file_path)
            
            # Vérification si déjà traité
            if self._is_document_processed(file_path.name, file_hash):
                return {
                    'success': True, 
                    'skipped': True, 
                    'message': 'Document déjà traité (hash identique)'
                }
            
            # Détection du type de fichier
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_types:
                return {
                    'success': False, 
                    'message': f'Type de fichier non supporté: {file_extension}'
                }
            
            logger.info(f"📄 Traitement: {file_path.name}")
            
            # Extraction du contenu
            processor_func = self.supported_types[file_extension]
            content, metadata = processor_func(file_path)
            
            if not content or len(content.strip()) < 10:
                return {'success': False, 'message': 'Contenu vide ou trop court'}
            
            # Découpage en chunks
            chunks = self._create_chunks(content, metadata)
            
            # Enregistrement en base
            document_id = self._save_document_record(
                file_path, file_hash, len(chunks), metadata
            )
            
            # Indexation dans le RAG
            chunks_created = 0
            for i, chunk in enumerate(chunks):
                try:
                    # Ajout du contexte document dans les métadonnées du chunk
                    chunk_metadata = {
                        'source_document': file_path.name,
                        'document_type': file_extension[1:],  # Sans le point
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'document_metadata': metadata,
                        'processing_date': datetime.now().isoformat()
                    }
                    
                    # Ajout au RAG
                    self.rag_service.add_knowledge_item(
                        content=chunk,
                        metadata=chunk_metadata
                    )
                    
                    # Sauvegarde du chunk en base
                    self._save_chunk_record(document_id, i, chunk, chunk_metadata)
                    chunks_created += 1
                    
                except Exception as e:
                    logger.warning(f"⚠️ Erreur indexation chunk {i}: {e}")
            
            # Déplacement vers le dossier processed
            processed_path = self.processed_dir / file_path.name
            if not processed_path.exists():
                file_path.rename(processed_path)
            
            result = {
                'success': True,
                'chunks_created': chunks_created,
                'message': f'Document traité: {chunks_created} chunks créés'
            }
            
            logger.info(f"✅ {file_path.name}: {chunks_created} chunks indexés")
            return result
            
        except Exception as e:
            logger.error(f"❌ Erreur traitement {file_path.name}: {e}")
            
            # Déplacement vers le dossier failed
            try:
                failed_path = self.failed_dir / file_path.name
                if not failed_path.exists():
                    file_path.rename(failed_path)
            except:
                pass
            
            return {'success': False, 'message': str(e)}
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calcule le hash MD5 d'un fichier"""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception:
            return ""
    
    def _is_document_processed(self, filename: str, file_hash: str) -> bool:
        """Vérifie si un document a déjà été traité"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    "SELECT id FROM processed_documents WHERE filename = ? AND file_hash = ?",
                    (filename, file_hash)
                )
                return cursor.fetchone() is not None
        except Exception:
            return False
    
    def _save_document_record(self, file_path: Path, file_hash: str, 
                            chunks_count: int, metadata: Dict) -> int:
        """Sauvegarde l'enregistrement du document"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    INSERT OR REPLACE INTO processed_documents 
                    (filename, file_path, file_hash, file_size, file_type, 
                     chunks_created, processing_status, metadata)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    file_path.name,
                    str(file_path),
                    file_hash,
                    file_path.stat().st_size,
                    file_path.suffix,
                    chunks_count,
                    'success',
                    json.dumps(metadata, ensure_ascii=False)
                ))
                conn.commit()
                return cursor.lastrowid
        except Exception as e:
            logger.error(f"Erreur sauvegarde document: {e}")
            return 0
    
    def _save_chunk_record(self, document_id: int, chunk_index: int, 
                          content: str, metadata: Dict):
        """Sauvegarde l'enregistrement d'un chunk"""
        try:
            content_hash = hashlib.md5(content.encode()).hexdigest()
            
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("""
                    INSERT OR IGNORE INTO document_chunks 
                    (document_id, chunk_index, content, content_hash, metadata)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    document_id,
                    chunk_index,
                    content,
                    content_hash,
                    json.dumps(metadata, ensure_ascii=False)
                ))
                conn.commit()
        except Exception as e:
            logger.warning(f"Erreur sauvegarde chunk: {e}")
    
    def _create_chunks(self, content: str, metadata: Dict) -> List[str]:
        """Découpe le contenu en chunks optimaux pour le RAG"""
        if len(content) <= self.chunk_config['max_chunk_size']:
            return [content.strip()]
        
        chunks = []
        sentences = self._split_into_sentences(content)
        
        current_chunk = ""
        
        for sentence in sentences:
            # Si ajouter cette phrase dépasse la taille max
            if len(current_chunk) + len(sentence) > self.chunk_config['max_chunk_size']:
                if len(current_chunk) >= self.chunk_config['min_chunk_size']:
                    chunks.append(current_chunk.strip())
                    
                    # Overlap: garder les derniers mots pour le contexte
                    words = current_chunk.split()
                    overlap_words = words[-self.chunk_config['overlap_size']//10:]
                    current_chunk = " ".join(overlap_words) + " " + sentence
                else:
                    current_chunk += " " + sentence
            else:
                current_chunk += " " + sentence
        
        # Ajouter le dernier chunk
        if len(current_chunk.strip()) >= self.chunk_config['min_chunk_size']:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Découpe le texte en phrases"""
        import re
        
        # Pattern pour détecter la fin des phrases
        sentence_endings = re.compile(r'[.!?]+\s+')
        sentences = sentence_endings.split(text)
        
        # Nettoyer et filtrer
        clean_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Phrases minimales
                clean_sentences.append(sentence)
        
        return clean_sentences
    
    # =======================================================================
    # PROCESSEURS PAR TYPE DE FICHIER
    # =======================================================================
    
    def _process_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier PDF"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 et pdfplumber requis pour les PDF")
        
        content = ""
        metadata = {
            'pages': 0,
            'extraction_method': 'pdfplumber',
            'title': '',
            'author': ''
        }
        
        try:
            # Essai avec pdfplumber (meilleur pour la mise en forme)
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n\n--- Page {i+1} ---\n"
                        content += page_text
        
        except Exception as e:
            logger.warning(f"Erreur pdfplumber, essai PyPDF2: {e}")
            
            # Fallback avec PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    metadata['extraction_method'] = 'PyPDF2'
                    
                    # Métadonnées du PDF
                    if pdf_reader.metadata:
                        metadata['title'] = str(pdf_reader.metadata.get('/Title', ''))
                        metadata['author'] = str(pdf_reader.metadata.get('/Author', ''))
                    
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n\n--- Page {i+1} ---\n"
                            content += page_text
            
            except Exception as e2:
                raise Exception(f"Erreur extraction PDF: {e2}")
        
        return content.strip(), metadata
    
    def _process_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier texte (TXT, MD)"""
        try:
            # Détection de l'encodage
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = ""
            used_encoding = 'utf-8'
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                raise Exception("Impossible de décoder le fichier texte")
            
            metadata = {
                'encoding': used_encoding,
                'lines': len(content.split('\n')),
                'words': len(content.split()),
                'file_type': 'markdown' if file_path.suffix == '.md' else 'text'
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Erreur lecture fichier texte: {e}")
    
    def _process_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier Word (DOCX)"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx requis pour les fichiers Word")
        
        try:
            doc = DocxDocument(file_path)
            
            content = ""
            paragraphs = 0
            
            # Extraction du texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
                    paragraphs += 1
            
            # Extraction du texte des tableaux
            tables_content = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_content += row_text + "\n"
            
            if tables_content:
                content += "\n\n--- Tableaux ---\n" + tables_content
            
            metadata = {
                'paragraphs': paragraphs,
                'tables': len(doc.tables),
                'core_properties': {}
            }
            
            # Propriétés du document
            try:
                props = doc.core_properties
                metadata['core_properties'] = {
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else ''
                }
            except:
                pass
            
            return content.strip(), metadata
            
        except Exception as e:
            raise Exception(f"Erreur traitement DOCX: {e}")
    
    def _process_excel(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier Excel"""
        try:
            # Lecture avec pandas
            excel_data = pd.read_excel(file_path, sheet_name=None)  # Toutes les feuilles
            
            content = ""
            metadata = {
                'sheets': list(excel_data.keys()),
                'total_sheets': len(excel_data),
                'total_rows': 0,
                'total_columns': 0
            }
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                content += f"\n\n--- Feuille: {sheet_name} ---\n"
                
                # Statistiques de la feuille
                rows, cols = df.shape
                metadata['total_rows'] += rows
                metadata['total_columns'] = max(metadata['total_columns'], cols)
                
                content += f"Dimensions: {rows} lignes × {cols} colonnes\n\n"
                
                # Colonnes
                content += "Colonnes: " + ", ".join(df.columns.astype(str)) + "\n\n"
                
                # Échantillon de données (premières lignes)
                sample_size = min(10, len(df))
                if sample_size > 0:
                    content += "Échantillon de données:\n"
                    content += df.head(sample_size).to_string(index=False) + "\n"
                
                # Statistiques descriptives pour les colonnes numériques
                numeric_cols = df.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 0:
                    content += "\nStatistiques numériques:\n"
                    content += df[numeric_cols].describe().to_string() + "\n"
            
            return content.strip(), metadata
            
        except Exception as e:
            raise Exception(f"Erreur traitement Excel: {e}")
    
    def _process_csv(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier CSV"""
        try:
            # Détection du séparateur et lecture
            with open(file_path, 'r', encoding='utf-8') as f:
                first_line = f.readline()
            
            separator = ',' if first_line.count(',') > first_line.count(';') else ';'
            
            df = pd.read_csv(file_path, sep=separator, encoding='utf-8')
            
            rows, cols = df.shape
            
            content = f"Fichier CSV: {file_path.name}\n"
            content += f"Dimensions: {rows} lignes × {cols} colonnes\n\n"
            
            # Colonnes
            content += "Colonnes:\n"
            for i, col in enumerate(df.columns):
                dtype = str(df[col].dtype)
                null_count = df[col].isnull().sum()
                content += f"  {i+1}. {col} ({dtype}) - {null_count} valeurs manquantes\n"
            
            # Échantillon de données
            sample_size = min(20, len(df))
            if sample_size > 0:
                content += f"\nÉchantillon ({sample_size} premières lignes):\n"
                content += df.head(sample_size).to_string(index=False) + "\n"
            
            # Statistiques descriptives
            content += "\nStatistiques descriptives:\n"
            content += df.describe(include='all').to_string() + "\n"
            
            metadata = {
                'rows': rows,
                'columns': cols,
                'separator': separator,
                'column_names': df.columns.tolist(),
                'dtypes': df.dtypes.astype(str).to_dict(),
                'missing_values': df.isnull().sum().to_dict()
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Erreur traitement CSV: {e}")
    
    def _process_json(self, file_path: Path) -> Tuple[str, Dict]:
        """Traite un fichier JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            content = f"Fichier JSON: {file_path.name}\n\n"
            
            # Analyse de la structure
            def analyze_json_structure(obj, level=0):
                indent = "  " * level
                result = ""
                
                if isinstance(obj, dict):
                    result += f"{indent}Objet ({len(obj)} clés):\n"
                    for key, value in obj.items():
                        result += f"{indent}  {key}: {type(value).__name__}\n"
                        if level < 2:  # Limiter la profondeur
                            result += analyze_json_structure(value, level + 1)
                elif isinstance(obj, list):
                    result += f"{indent}Liste ({len(obj)} éléments)\n"
                    if obj and level < 2:
                        result += f"{indent}  Type d'éléments: {type(obj[0]).__name__}\n"
                        result += analyze_json_structure(obj[0], level + 1)
                
                return result
            
            content += "Structure:\n"
            content += analyze_json_structure(data)
            
            # Contenu sérialisé (limité)
            json_str = json.dumps(data, indent=2, ensure_ascii=False)
            if len(json_str) > 5000:
                json_str = json_str[:5000] + "\n... (contenu tronqué)"
            
            content += f"\nContenu:\n{json_str}"
            
            metadata = {
                'json_type': type(data).__name__,
                'size_chars': len(json_str),
                'structure_depth': self._calculate_json_depth(data)
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Erreur traitement JSON: {e}")
    
    def _calculate_json_depth(self, obj, current_depth=0):
        """Calcule la profondeur d'un objet JSON"""
        if isinstance(obj, dict):
            return max([self._calculate_json_depth(value, current_depth + 1) 
                       for value in obj.values()] + [current_depth])
        elif isinstance(obj, list) and obj:
            return max([self._calculate_json_depth(item, current_depth + 1) 
                       for item in obj] + [current_depth])
        else:
            return current_depth
    
    # =======================================================================
    # GESTION ET STATISTIQUES
    # =======================================================================
    
    def get_processing_stats(self) -> Dict:
        """Retourne les statistiques de traitement"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # Documents traités
                cursor = conn.execute("""
                    SELECT COUNT(*) as total,
                           SUM(chunks_created) as total_chunks,
                           AVG(chunks_created) as avg_chunks_per_doc
                    FROM processed_documents 
                    WHERE processing_status = 'success'
                """)
                doc_stats = cursor.fetchone()
                
                # Par type de fichier
                cursor = conn.execute("""
                    SELECT file_type, COUNT(*) as count, SUM(chunks_created) as chunks
                    FROM processed_documents 
                    WHERE processing_status = 'success'
                    GROUP BY file_type
                """)
                type_stats = cursor.fetchall()
                
                # Traitement récent
                cursor = conn.execute("""
                    SELECT COUNT(*) 
                    FROM processed_documents 
                    WHERE processing_date > datetime('now', '-7 days')
                """)
                recent_count = cursor.fetchone()[0]
                
                return {
                    'total_documents': doc_stats[0] or 0,
                    'total_chunks': doc_stats[1] or 0,
                    'avg_chunks_per_document': round(doc_stats[2] or 0, 1),
                    'documents_by_type': {
                        row[0]: {'count': row[1], 'chunks': row[2]} 
                        for row in type_stats
                    },
                    'documents_last_week': recent_count,
                    'supported_formats': list(self.supported_types.keys())
                }
                
        except Exception as e:
            logger.error(f"Erreur stats processing: {e}")
            return {}
    
    def list_processed_documents(self) -> List[Dict]:
        """Liste des documents traités"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    SELECT filename, file_type, processing_date, chunks_created, 
                           file_size, processing_status
                    FROM processed_documents
                    ORDER BY processing_date DESC
                """)
                
                results = []
                for row in cursor.fetchall():
                    results.append({
                        'filename': row[0],
                        'file_type': row[1],
                        'processing_date': row[2],
                        'chunks_created': row[3],
                        'file_size_mb': round(row[4] / (1024*1024), 2),
                        'status': row[5]
                    })
                
                return results
                
        except Exception as e:
            logger.error(f"Erreur liste documents: {e}")
            return []
    
    def remove_document_from_rag(self, filename: str) -> bool:
        """Supprime un document du RAG"""
        try:
            # TODO: Implémenter la suppression dans le RAG service
            # En attendant, marquer comme supprimé en base
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    "UPDATE processed_documents SET processing_status = 'deleted' WHERE filename = ?",
                    (filename,)
                )
                conn.commit()
                return True
                
        except Exception as e:
            logger.error(f"Erreur suppression document: {e}")
            return False
    
    def clear_all_documents(self) -> bool:
        """Vide toute la base de documents"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("DELETE FROM document_chunks")
                conn.execute("DELETE FROM processed_documents")
                conn.commit()
                
            logger.info("🗑️ Base de documents vidée")
            return True
            
        except Exception as e:
            logger.error(f"Erreur vidage base: {e}")
            return False


# ==============================================================================
# SERVICE DE SURVEILLANCE DES DOSSIERS
# ==============================================================================

class DocumentWatcher:
    """Surveillant de dossier pour traitement automatique"""
    
    def __init__(self, document_processor: DocumentProcessor, watch_interval: int = 60):
        """
        Initialise le surveillant de dossier
        
        Args:
            document_processor: Instance du processeur
            watch_interval: Intervalle de surveillance en secondes
        """
        self.processor = document_processor
        self.watch_interval = watch_interval
        self.is_watching = False
        self.last_scan = None
        
    def start_watching(self):
        """Démarre la surveillance du dossier"""
        import threading
        import time
        
        self.is_watching = True
        
        def watch_loop():
            logger.info(f"👁️ Surveillance dossier démarrée: {self.processor.documents_dir}")
            
            while self.is_watching:
                try:
                    # Scan et traitement
                    results = self.processor.process_documents_directory()
                    
                    if results['processed'] > 0:
                        logger.info(f"📁 Nouveaux documents traités: {results['processed']}")
                    
                    self.last_scan = datetime.now()
                    
                    # Attente avant le prochain scan
                    time.sleep(self.watch_interval)
                    
                except Exception as e:
                    logger.error(f"Erreur surveillance: {e}")
                    time.sleep(self.watch_interval)
        
        # Lancement en thread séparé
        watch_thread = threading.Thread(target=watch_loop, daemon=True)
        watch_thread.start()
        
        return watch_thread
    
    def stop_watching(self):
        """Arrête la surveillance"""
        self.is_watching = False
        logger.info("⏹️ Surveillance dossier arrêtée")


# ==============================================================================
# API REST POUR L'INTERFACE WEB
# ==============================================================================

def create_document_api_routes(app, document_processor):
    """Crée les routes API pour la gestion des documents"""
    
    @app.route('/api/documents/upload', methods=['POST'])
    def upload_documents():
        """Upload de documents via l'interface web"""
        try:
            from flask import request, jsonify
            import werkzeug
            
            if 'files' not in request.files:
                return jsonify({'success': False, 'error': 'Aucun fichier'}), 400
            
            files = request.files.getlist('files')
            results = []
            
            for file in files:
                if file.filename == '':
                    continue
                
                # Sauvegarde temporaire
                temp_path = document_processor.documents_dir / file.filename
                file.save(temp_path)
                
                # Traitement
                result = document_processor.process_single_document(temp_path)
                result['filename'] = file.filename
                results.append(result)
            
            successful = sum(1 for r in results if r['success'])
            
            return jsonify({
                'success': True,
                'message': f'{successful}/{len(results)} fichiers traités',
                'results': results
            })
            
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/api/documents/stats')
    def documents_stats():
        """Statistiques des documents"""
        try:
            stats = document_processor.get_processing_stats()
            return jsonify({'success': True, 'stats': stats})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/api/documents/list')
    def documents_list():
        """Liste des documents traités"""
        try:
            documents = document_processor.list_processed_documents()
            return jsonify({'success': True, 'documents': documents})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/api/documents/process-all', methods=['POST'])
    def process_all_documents():
        """Traite tous les documents du dossier"""
        try:
            results = document_processor.process_documents_directory()
            return jsonify({'success': True, 'results': results})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/api/documents/remove/<filename>', methods=['DELETE'])
    def remove_document(filename):
        """Supprime un document du RAG"""
        try:
            success = document_processor.remove_document_from_rag(filename)
            if success:
                return jsonify({'success': True, 'message': 'Document supprimé'})
            else:
                return jsonify({'success': False, 'error': 'Échec suppression'}), 500
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500


# ==============================================================================
# UTILITAIRES ET TESTS
# ==============================================================================

def install_optional_dependencies():
    """Installe les dépendances optionnelles pour le traitement de documents"""
    import subprocess
    import sys
    
    optional_packages = [
        'PyPDF2',           # PDF
        'pdfplumber',       # PDF amélioré
        'python-docx',      # Word
        'openpyxl',         # Excel
        'python-dotenv',    # Variables d'environnement
        'watchdog'          # Surveillance de fichiers
    ]
    
    print("📦 Installation des dépendances pour le traitement de documents...")
    
    for package in optional_packages:
        try:
            subprocess.check_call([
                sys.executable, '-m', 'pip', 'install', package, '--quiet'
            ])
            print(f"✅ {package}")
        except subprocess.CalledProcessError:
            print(f"❌ Échec {package}")


def create_example_documents(documents_dir: str = "documents"):
    """Crée des documents d'exemple pour tester le système"""
    docs_path = Path(documents_dir)
    docs_path.mkdir(exist_ok=True)
    
    # 1. Fichier texte sur l'énergie en Malaysia
    energy_guide = docs_path / "malaysia_energy_guide.txt"
    energy_guide.write_text("""
Guide de l'Énergie en Malaysia

La Malaysia est un pays d'Asie du Sud-Est avec une consommation énergétique en croissance rapide.
Le pays dispose de ressources énergétiques diversifiées incluant le pétrole, le gaz naturel, 
le charbon et un potentiel significatif pour les énergies renouvelables.

CONSOMMATION ÉLECTRIQUE:
- Secteur résidentiel: 25% de la consommation totale
- Secteur industriel: 40% de la consommation totale  
- Secteur commercial: 30% de la consommation totale
- Transport: 5% de la consommation totale

TYPES DE BÂTIMENTS:
1. Résidentiels: Maisons individuelles, appartements, condominiums
2. Commerciaux: Centres commerciaux, bureaux, hôtels, restaurants
3. Industriels: Usines, entrepôts, centres de production
4. Publics: Hôpitaux, écoles, bâtiments gouvernementaux

ZONES GÉOGRAPHIQUES:
- Kuala Lumpur: Centre économique principal
- Selangor: Zone industrielle importante
- Johor: Frontière avec Singapour, forte activité économique
- Penang: Centre technologique et touristique
- Sabah & Sarawak: Régions orientales avec ressources naturelles

TENDANCES ÉNERGÉTIQUES:
- Croissance de 3-5% par an de la demande électrique
- Initiatives gouvernementales pour l'efficacité énergétique
- Développement des énergies solaires et éoliennes
- Smart grids et digitalisation du réseau électrique

DÉFIS:
- Réduction des émissions de CO2
- Diversification du mix énergétique
- Amélioration de l'efficacité énergétique des bâtiments
- Intégration des énergies renouvelables
""", encoding='utf-8')
    
    # 2. Fichier JSON avec données techniques
    tech_specs = docs_path / "technical_specifications.json"
    tech_specs.write_text(json.dumps({
        "building_efficiency_standards": {
            "residential": {
                "excellent": "< 50 kWh/m²/an",
                "good": "50-100 kWh/m²/an", 
                "average": "100-150 kWh/m²/an",
                "poor": "> 150 kWh/m²/an"
            },
            "commercial": {
                "excellent": "< 100 kWh/m²/an",
                "good": "100-200 kWh/m²/an",
                "average": "200-300 kWh/m²/an", 
                "poor": "> 300 kWh/m²/an"
            }
        },
        "climate_zones": {
            "tropical_humid": {
                "description": "Climat tropical humide de Malaysia",
                "temperature_range": "24-35°C",
                "humidity": "60-90%",
                "cooling_degree_days": "3000-4000"
            }
        },
        "equipment_specifications": {
            "air_conditioning": {
                "cop_minimum": 3.0,
                "recommended_cop": 4.5,
                "sizing_guidelines": "150-200 W/m²"
            },
            "lighting": {
                "led_efficiency": "100-150 lm/W",
                "recommended_levels": {
                    "office": "300-500 lux",
                    "residential": "100-300 lux",
                    "retail": "500-1000 lux"
                }
            }
        }
    }, indent=2, ensure_ascii=False), encoding='utf-8')
    
    # 3. Fichier CSV avec données de référence
    reference_data = docs_path / "reference_buildings.csv"
    reference_data.write_text("""building_type,typical_consumption_kwh_m2,peak_demand_w_m2,operating_hours,efficiency_class
residential,120,45,16,B
commercial_office,180,55,12,C
retail_shop,220,65,14,C
restaurant,350,85,16,D
hotel,200,60,24,B
hospital,400,95,24,E
school,150,50,10,B
warehouse,80,35,8,A
factory_light,250,75,16,C
factory_heavy,450,120,20,D
""", encoding='utf-8')
    
    print(f"📝 Documents d'exemple créés dans {docs_path}/")
    return docs_path


def test_document_processor():
    """Test du processeur de documents"""
    
    # Création des documents d'exemple
    docs_dir = create_example_documents("test_documents")
    
    # Création d'un RAG service de test
    from dashboard.services.rag_service import RAGService
    rag_service = RAGService("test_rag.db")
    
    # Initialisation du processeur
    processor = DocumentProcessor(rag_service, "test_documents")
    
    print("🧪 Test DocumentProcessor:")
    
    # Test de traitement
    results = processor.process_documents_directory()
    
    print(f"📊 Résultats:")
    print(f"   Fichiers traités: {results['processed']}")
    print(f"   Chunks créés: {results['new_chunks']}")
    print(f"   Échecs: {results['failed']}")
    
    # Statistiques
    stats = processor.get_processing_stats()
    print(f"📈 Statistiques:")
    print(f"   Documents total: {stats['total_documents']}")
    print(f"   Chunks total: {stats['total_chunks']}")
    print(f"   Par type: {stats['documents_by_type']}")
    
    # Liste des documents
    docs = processor.list_processed_documents()
    print(f"📋 Documents traités:")
    for doc in docs:
        print(f"   • {doc['filename']} ({doc['file_type']}) - {doc['chunks_created']} chunks")
    
    return processor


if __name__ == '__main__':
    # Vérification des dépendances
    missing_deps = []
    
    if not PDF_AVAILABLE:
        missing_deps.append("PyPDF2, pdfplumber (pour PDF)")
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx (pour Word)")
    if not EXCEL_AVAILABLE:
        missing_deps.append("openpyxl (pour Excel)")
    
    if missing_deps:
        print("⚠️ Dépendances manquantes:")
        for dep in missing_deps:
            print(f"   • {dep}")
        print("\n💡 Installez avec: pip install PyPDF2 pdfplumber python-docx openpyxl")
        print("   Ou utilisez: python document_processor.py --install-deps")
    else:
        print("✅ Toutes les dépendances sont installées")
    
    # Test si demandé
    import sys
    if '--test' in sys.argv:
        test_document_processor()
    elif '--install-deps' in sys.argv:
        install_optional_dependencies()
    elif '--create-examples' in sys.argv:
        create_example_documents()
        print("✅ Documents d'exemple créés")